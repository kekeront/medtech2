"""FastAPI application implementing the TZ 4.5 search API plus upload / review endpoints.

Run: uvicorn app.api:app --reload
Docs: /docs (Swagger) and /openapi.json
"""

from __future__ import annotations

import tempfile
import time
import uuid
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from fastapi import (
    BackgroundTasks,
    Depends,
    FastAPI,
    File,
    HTTPException,
    Query,
    UploadFile,
)
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy import delete as sa_delete
from sqlalchemy import func, or_, select
from sqlalchemy.orm import Session

from . import catalogue
from .archive import SUPPORTED_SUFFIXES, ensure_parseable, ingest_zip
from .config import ORIGINALS_DIR
from .db import get_session, init_db
from .extract_api import router as extract_router
from .fx_api import router as fx_router
from .models import Partner, PriceDocument, PriceItem, Service
from .normalize import _key
from .ocr_api import router as ocr_router
from .pipeline import ingest_file
from .schemas import (
    DocumentOut,
    MatchIn,
    NewItemIn,
    PartnerOut,
    PartnerPriceOut,
    PatchItemIn,
    PriceItemOut,
    ServiceIn,
    ServiceOut,
    VerifyIn,
)

app = FastAPI(
    title="MedArchive API",
    version="1.0",
    description="Partner clinics, normalized services and prices extracted from price-list archives.",
)

STATIC_DIR = Path(__file__).parent / "static"


app.include_router(ocr_router)
app.include_router(fx_router)
app.include_router(extract_router)


@app.on_event("startup")
def _startup() -> None:
    init_db()


# ----------------------------------------------------------------------------- services
@app.get("/services", response_model=list[ServiceOut], tags=["services"])
def list_services(
    category: str | None = None,
    q: str | None = None,
    limit: int = Query(100, le=1000),
    offset: int = 0,
    db: Session = Depends(get_session),
):
    stmt = select(Service).where(Service.is_active.is_(True))
    if category:
        stmt = stmt.where(Service.category == category)
    if q:
        stmt = stmt.where(Service.service_name.ilike(f"%{q}%"))
    stmt = stmt.order_by(Service.service_name).limit(limit).offset(offset)
    return db.scalars(stmt).all()


@app.post("/services", response_model=ServiceOut, tags=["services"], status_code=201)
def create_service(body: ServiceIn, db: Session = Depends(get_session)):
    # Dedup on the normalized key so two spellings of one service (e.g. "УЗИ" vs
    # "Ультразвуковое исследование") don't create colliding catalogue rows.
    key = _key(body.service_name)
    existing = next(
        (s for s in db.scalars(select(Service)) if _key(s.service_name) == key), None
    )
    if existing is not None:
        existing.synonyms = list(
            dict.fromkeys((existing.synonyms or []) + (body.synonyms or []))
        )
        existing.category = existing.category or body.category
        existing.icd_code = existing.icd_code or body.icd_code
        db.commit()
        return existing
    svc = Service(
        service_name=body.service_name,
        category=body.category,
        synonyms=body.synonyms,
        icd_code=body.icd_code,
    )
    db.add(svc)
    db.commit()
    return svc


@app.post("/services/import", tags=["services"])
async def import_catalogue(
    file: UploadFile = File(...),
    rematch_after: bool = True,
    db: Session = Depends(get_session),
):
    """Load the target service catalogue from an uploaded XLSX/JSON file (TZ 2.2)."""
    name = file.filename or "catalogue"
    suffix = Path(name).suffix.lower()
    if suffix not in {".xlsx", ".xls", ".json"}:
        raise HTTPException(415, f"unsupported catalogue format: {suffix}")
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        created = catalogue.load_catalogue_from_file(db, tmp_path)
        if rematch_after:
            stats = catalogue.rematch(db)  # commits load + matches atomically
        else:
            stats = {}
            db.commit()  # load alone now only flushes — persist it
    except Exception as e:  # noqa: BLE001
        db.rollback()
        raise HTTPException(
            422, f"catalogue import failed: {type(e).__name__}: {e}"
        ) from e
    finally:
        Path(tmp_path).unlink(missing_ok=True)
    return {"services_created": created, "rematch": stats}


@app.post("/catalogue/bootstrap", tags=["services"])
def bootstrap_catalogue(min_count: int = 2, db: Session = Depends(get_session)):
    """Synthesize a catalogue from already-extracted positions when none was provided (TZ 7)."""
    created = catalogue.bootstrap_catalogue(db, min_count=min_count)
    return {"services_created": created, "rematch": catalogue.rematch(db)}


@app.post("/rematch", tags=["services"])
def rematch(only_unmatched: bool = True, db: Session = Depends(get_session)):
    """Re-run catalogue matching over price items (TZ 4.3)."""
    return catalogue.rematch(db, only_unmatched=only_unmatched)


@app.get(
    "/services/{service_id}/partners",
    response_model=list[PartnerPriceOut],
    tags=["services"],
)
def service_partners(service_id: uuid.UUID, db: Session = Depends(get_session)):
    if not db.get(Service, service_id):
        raise HTTPException(404, "service not found")
    rows = db.scalars(
        select(PriceItem)
        .where(PriceItem.service_id == service_id, PriceItem.is_active.is_(True))
        .order_by(PriceItem.price_resident_kzt.asc().nulls_last())
    ).all()
    return [_partner_price(db, it) for it in rows]


# ----------------------------------------------------------------------------- partners
@app.get("/partners", response_model=list[PartnerOut], tags=["partners"])
def list_partners(
    city: str | None = None,
    is_active: bool | None = None,
    limit: int = Query(100, le=1000),
    offset: int = 0,
    db: Session = Depends(get_session),
):
    stmt = select(Partner)
    if city:
        stmt = stmt.where(Partner.city.ilike(f"%{city}%"))
    if is_active is not None:
        stmt = stmt.where(Partner.is_active.is_(is_active))
    stmt = stmt.order_by(Partner.name).limit(limit).offset(offset)
    return db.scalars(stmt).all()


@app.get("/partners/{partner_id}", response_model=PartnerOut, tags=["partners"])
def get_partner(partner_id: uuid.UUID, db: Session = Depends(get_session)):
    partner = db.get(Partner, partner_id)
    if not partner:
        raise HTTPException(404, "partner not found")
    return partner


def _price_key(it: PriceItem) -> tuple:
    """Identity for preview aggregation: name + every price field (resident, non-res,
    extra tiers, currency). Two rows match only when EXACTLY identical here — never
    fuses different services or different prices."""

    def n(v) -> float | None:
        return round(float(v), 2) if v is not None else None

    tiers = it.price_extra_tiers or {}
    return (
        (it.service_name_raw or "").strip(),
        n(it.price_resident_kzt),
        n(it.price_nonresident_kzt),
        tuple(sorted((k, n(v)) for k, v in tiers.items())),
        it.currency_original,
    )


def _aggregate_identical(items: list[PriceItem]) -> list[PriceItemOut]:
    """Collapse rows that are exactly identical in name+price into one preview row.

    Storage keeps every PriceItem — two codes (e.g. U2.1.5 / U2.2.3) may diverge in a
    future price-list version, so they must stay individually tracked. Only the preview
    merges true name+price twins; first-occurrence order (section, name) is preserved.
    """
    by_key: dict[tuple, PriceItemOut] = {}
    out: list[PriceItemOut] = []
    for it in items:
        key = _price_key(it)
        first = by_key.get(key)
        if first is None:
            row = PriceItemOut.model_validate(it)
            row.merged_codes = (
                [it.service_code_source] if it.service_code_source else []
            )
            by_key[key] = row
            out.append(row)
        else:
            first.merged_count += 1
            code = it.service_code_source
            if code and code not in (first.merged_codes or []):
                first.merged_codes.append(code)
    for row in out:  # tidy: no badge data when the row stands for a single source row
        if row.merged_count == 1 or not row.merged_codes:
            row.merged_codes = None
    return out


@app.get(
    "/partners/{partner_id}/services",
    response_model=list[PriceItemOut],
    tags=["partners"],
)
def partner_services(
    partner_id: uuid.UUID,
    active_only: bool = True,
    aggregate: bool = True,
    limit: int = Query(2000, le=10000),
    offset: int = 0,
    db: Session = Depends(get_session),
):
    """Partner price list (the DB preview). By default, rows that are exactly identical
    in name and price are aggregated into one displayed service (`merged_count` > 1);
    pass `aggregate=false` for the raw, per-row list."""
    if not db.get(Partner, partner_id):
        raise HTTPException(404, "partner not found")
    stmt = select(PriceItem).where(PriceItem.partner_id == partner_id)
    if active_only:
        stmt = stmt.where(PriceItem.is_active.is_(True))
    stmt = (
        stmt.order_by(PriceItem.section, PriceItem.service_name_raw)
        .limit(limit)
        .offset(offset)
    )
    items = db.scalars(stmt).all()
    return _aggregate_identical(items) if aggregate else items


# ----------------------------------------------------------------------------- search
@app.get("/search", tags=["search"])
def search(
    q: str = Query(..., min_length=2),
    limit: int = Query(50, le=500),
    db: Session = Depends(get_session),
):
    like = f"%{q}%"
    # Russian-aware term: homoglyphs folded + abbreviations expanded (УЗИ↔ультразвук…).
    norm = f"%{_key(q)}%"
    partners = db.scalars(
        select(Partner).where(Partner.name.ilike(like)).limit(limit)
    ).all()
    services = db.scalars(
        select(Service)
        .where(Service.service_name.ilike(like), Service.is_active.is_(True))
        .limit(limit)
    ).all()
    # Group same service together and order cheapest-first so an insurance operator can
    # compare which clinic offers the service at the best resident price.
    items = db.scalars(
        select(PriceItem)
        .where(
            or_(
                PriceItem.service_name_raw.ilike(like), PriceItem.name_norm.ilike(norm)
            ),
            PriceItem.is_active.is_(True),
        )
        .order_by(
            PriceItem.name_norm,
            PriceItem.price_resident_kzt.asc().nulls_last(),
        )
        .limit(limit)
    ).all()
    return {
        "partners": [PartnerOut.model_validate(p).model_dump() for p in partners],
        "services": [ServiceOut.model_validate(s).model_dump() for s in services],
        "price_items": _with_partner_names(db, items),
    }


def _with_partner_names(db: Session, items) -> list[dict]:
    """Serialize price items and attach the clinic name (operator context)."""
    names = {p.partner_id: p.name for p in db.scalars(select(Partner))}
    out = []
    for i in items:
        d = PriceItemOut.model_validate(i).model_dump()
        d["partner_name"] = names.get(i.partner_id)
        out.append(d)
    return out


# ----------------------------------------------------------------- operator / review (TZ 4.3-4.4)
@app.get("/unmatched", response_model=list[PriceItemOut], tags=["operator"])
def unmatched(
    limit: int = Query(200, le=2000),
    offset: int = 0,
    db: Session = Depends(get_session),
):
    stmt = (
        select(PriceItem)
        .where(PriceItem.service_id.is_(None), PriceItem.is_active.is_(True))
        .order_by(PriceItem.match_confidence.desc().nulls_last())
        .limit(limit)
        .offset(offset)
    )
    return _with_partner_names(db, db.scalars(stmt).all())


@app.get("/review", response_model=list[PriceItemOut], tags=["operator"])
def review_queue(
    limit: int = Query(200, le=2000),
    offset: int = 0,
    db: Session = Depends(get_session),
):
    stmt = (
        select(PriceItem)
        .where(PriceItem.needs_review.is_(True), PriceItem.is_verified.is_(False))
        .order_by(PriceItem.effective_date.desc().nulls_last())
        .limit(limit)
        .offset(offset)
    )
    return _with_partner_names(db, db.scalars(stmt).all())


@app.get("/flagged", response_model=list[PriceItemOut], tags=["operator"])
def flagged_queue(
    limit: int = Query(500, le=5000),
    offset: int = 0,
    db: Session = Depends(get_session),
):
    """Rows the one-pass LLM cleanup flagged as suspicious (price/tier/junk) for the Проверка
    page. These carry an 'LLM:' marker in review_reason and await human confirmation."""
    stmt = (
        select(PriceItem)
        .where(
            PriceItem.is_active.is_(True),
            PriceItem.is_verified.is_(False),
            PriceItem.review_reason.like("%LLM:%"),
        )
        .order_by(PriceItem.effective_date.desc().nulls_last())
        .limit(limit)
        .offset(offset)
    )
    return _with_partner_names(db, db.scalars(stmt).all())


@app.post("/match", response_model=list[PriceItemOut], tags=["operator"])
def match(body: MatchIn, db: Session = Depends(get_session)):
    if body.service_id is not None and not db.get(Service, body.service_id):
        raise HTTPException(404, "service not found")
    items = db.scalars(
        select(PriceItem).where(PriceItem.item_id.in_(body.item_ids))
    ).all()
    if not items:
        raise HTTPException(404, "no matching items")
    for it in items:
        it.service_id = body.service_id
        it.match_method = "manual" if body.service_id else None
        it.match_confidence = 1.0 if body.service_id else None
        if body.note:
            it.verification_note = body.note
    db.commit()
    return items


@app.patch("/items/{item_id}", response_model=PriceItemOut, tags=["operator"])
def patch_item(
    item_id: uuid.UUID, body: PatchItemIn, db: Session = Depends(get_session)
):
    """Partially update a price item's name, section, prices, or effective date."""
    it = db.get(PriceItem, item_id)
    if it is None:
        raise HTTPException(404, "item not found")
    for fld in body.model_fields_set:
        setattr(it, fld, getattr(body, fld))
    db.commit()
    db.refresh(it)
    return it


@app.delete("/items/{item_id}", status_code=204, tags=["operator"])
def delete_item(item_id: uuid.UUID, db: Session = Depends(get_session)):
    """Permanently delete a single price item."""
    it = db.get(PriceItem, item_id)
    if it is None:
        raise HTTPException(404, "item not found")
    db.delete(it)
    db.commit()


@app.get("/items/{item_id}/history", tags=["operator"])
def item_price_history(item_id: uuid.UUID, db: Session = Depends(get_session)):
    """Price dynamics for ONE service over the years: every dated price this partner has for
    the same service across all their documents, one point per effective date. Services are
    matched by the stable source code (e.g. U1.1) UNION normalized name — either is enough.
    The union matters when a clinic changes its code scheme between years (e.g. Клиника 1
    2024 vs 2026 share 0 codes but 96 names): code-only matching would never bridge the
    years, so the same service is also tracked by its normalized name."""
    it = db.get(PriceItem, item_id)
    if it is None:
        raise HTTPException(404, "item not found")

    conds = []
    parts = []
    if it.service_code_source:
        conds.append(PriceItem.service_code_source == it.service_code_source)
        parts.append("code")
    if it.name_norm:
        conds.append(PriceItem.name_norm == it.name_norm)
        parts.append("name")
    if not conds:
        conds.append(PriceItem.service_name_raw == it.service_name_raw)
        parts.append("name")
    match = or_(*conds)
    matched_by = "+".join(parts)

    rows = db.scalars(
        select(PriceItem)
        .where(
            PriceItem.partner_id == it.partner_id,
            PriceItem.is_active.is_(True),
            PriceItem.effective_date.isnot(None),
            match,
        )
        .order_by(PriceItem.effective_date)
    ).all()

    # Collapse duplicate re-ingests: one point per effective date, preferring a row that
    # actually has a resident price.
    by_date: dict[date, PriceItem] = {}
    for r in rows:
        cur = by_date.get(r.effective_date)
        if cur is None or (
            cur.price_resident_kzt is None and r.price_resident_kzt is not None
        ):
            by_date[r.effective_date] = r

    points = [
        {
            "effective_date": d.isoformat(),
            "year": d.year,
            "price_resident_kzt": float(r.price_resident_kzt)
            if r.price_resident_kzt is not None
            else None,
            "price_nonresident_kzt": float(r.price_nonresident_kzt)
            if r.price_nonresident_kzt is not None
            else None,
        }
        for d, r in sorted(by_date.items())
    ]
    return {
        "item_id": str(item_id),
        "service_name": it.service_name_raw,
        "service_code_source": it.service_code_source,
        "matched_by": matched_by,
        "points": points,
    }


@app.post("/items/{item_id}/verify", response_model=PriceItemOut, tags=["operator"])
def verify_item(item_id: uuid.UUID, body: VerifyIn, db: Session = Depends(get_session)):
    it = db.get(PriceItem, item_id)
    if not it:
        raise HTTPException(404, "item not found")
    it.is_verified = body.is_verified
    if body.is_verified:
        it.needs_review = False
    if body.price_resident_kzt is not None:
        it.price_resident_kzt = body.price_resident_kzt
    if body.price_nonresident_kzt is not None:
        it.price_nonresident_kzt = body.price_nonresident_kzt
    if body.note:
        it.verification_note = body.note
    db.commit()
    return it


# ----------------------------------------------------------------------------- documents / ingest
@app.get("/documents", response_model=list[DocumentOut], tags=["documents"])
def list_documents(
    status: str | None = None,
    limit: int = Query(200, le=2000),
    db: Session = Depends(get_session),
):
    stmt = select(PriceDocument).order_by(PriceDocument.parsed_at.desc())
    if status:
        stmt = stmt.where(PriceDocument.parse_status == status)
    return db.scalars(stmt.limit(limit)).all()


@app.delete("/documents/{doc_id}", tags=["documents"])
def delete_document(
    doc_id: uuid.UUID,
    purge_original: bool = False,
    db: Session = Depends(get_session),
):
    """Delete a document and ALL its extracted price items. The archived original file is
    kept by default (TZ 5: исходные файлы не удаляются); pass purge_original=true to also
    remove it from disk."""
    doc = db.get(PriceDocument, doc_id)
    if doc is None:
        raise HTTPException(404, "document not found")
    n_items = (
        db.scalar(
            select(func.count())
            .select_from(PriceItem)
            .where(PriceItem.doc_id == doc_id)
        )
        or 0
    )
    file_name = doc.file_name
    purged = False
    if purge_original and doc.content_sha256:
        path = ORIGINALS_DIR / f"{doc.content_sha256[:16]}__{Path(doc.file_name).name}"
        if path.exists():
            path.unlink()
            purged = True
    # Bulk-delete items first (avoids loading thousands of ORM rows for the cascade), then
    # the document; the now-empty relationship makes the ORM cascade a no-op.
    db.execute(sa_delete(PriceItem).where(PriceItem.doc_id == doc_id))
    db.delete(doc)
    db.commit()
    return {
        "deleted": str(doc_id),
        "file_name": file_name,
        "items_deleted": n_items,
        "original_purged": purged,
    }


def _safe_original_path(doc: PriceDocument) -> Path:
    """Build the ORIGINALS_DIR path for *doc* and verify it stays inside that directory."""
    if not doc.content_sha256:
        raise HTTPException(404, "original file not stored for this document")
    safe_name = Path(doc.file_name).name  # strip any directory components
    path = (ORIGINALS_DIR / f"{doc.content_sha256[:16]}__{safe_name}").resolve()
    if not str(path).startswith(str(ORIGINALS_DIR.resolve()) + "/"):
        raise HTTPException(404, "original file not found on disk")
    if not path.exists():
        raise HTTPException(404, "original file not found on disk")
    return path


_DOC_MEDIA_TYPES: dict[str, str] = {
    "pdf": "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
}


@app.get(
    "/documents/{doc_id}/items", response_model=list[PriceItemOut], tags=["documents"]
)
def document_items(doc_id: uuid.UUID, db: Session = Depends(get_session)):
    """All price items extracted from a single document."""
    if not db.get(PriceDocument, doc_id):
        raise HTTPException(404, "document not found")
    items = db.scalars(
        select(PriceItem)
        .where(PriceItem.doc_id == doc_id)
        .order_by(PriceItem.section, PriceItem.service_name_raw)
    ).all()
    return _aggregate_identical(list(items))


@app.post(
    "/documents/{doc_id}/items",
    response_model=PriceItemOut,
    tags=["documents"],
    status_code=201,
)
def create_document_item(
    doc_id: uuid.UUID, body: NewItemIn, db: Session = Depends(get_session)
):
    """Manually add a price item to a document (e.g. a row missed by OCR)."""
    doc = db.get(PriceDocument, doc_id)
    if doc is None:
        raise HTTPException(404, "document not found")
    item = PriceItem(
        doc_id=doc_id,
        partner_id=doc.partner_id,
        service_name_raw=body.service_name_raw,
        section=body.section,
        price_resident_kzt=body.price_resident_kzt,
        price_nonresident_kzt=body.price_nonresident_kzt,
        effective_date=body.effective_date,
        needs_review=True,
        review_reason="manually added",
        is_active=True,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.get("/documents/{doc_id}/file", tags=["documents"])
def document_file(doc_id: uuid.UUID, db: Session = Depends(get_session)):
    """Serve the original uploaded file for inline preview or download."""
    doc = db.get(PriceDocument, doc_id)
    if doc is None:
        raise HTTPException(404, "document not found")
    path = _safe_original_path(doc)
    media_type = _DOC_MEDIA_TYPES.get(
        doc.file_format.lower(), "application/octet-stream"
    )
    return FileResponse(str(path), filename=doc.file_name, media_type=media_type)


_PREVIEW_CSS = """
<style>
body{font-family:system-ui,sans-serif;font-size:13px;margin:0;padding:12px;color:#111}
table{border-collapse:collapse;width:100%;table-layout:auto}
th{background:#f5f5f5;border:1px solid #ddd;padding:6px 10px;text-align:left;white-space:nowrap;position:sticky;top:0}
td{border:1px solid #ddd;padding:5px 10px;white-space:nowrap;max-width:320px;overflow:hidden;text-overflow:ellipsis}
tr:nth-child(even){background:#fafafa}
</style>
"""

# Hardening for user-content previews: the table HTML is built from an UPLOADED file and
# rendered same-origin in an <iframe>. Cell/sheet text is already escaped (pandas to_html
# escape=True + html.escape), but as defence-in-depth the CSP blocks all scripts (the only XSS
# vector) while still allowing the inline styles and sheet-tab navigation the preview needs;
# nosniff stops the browser re-interpreting the body as another content type.
_PREVIEW_HEADERS = {
    "Content-Security-Policy": "default-src 'none'; style-src 'unsafe-inline'",
    "X-Content-Type-Options": "nosniff",
}


def _preview_html(body: str) -> HTMLResponse:
    return HTMLResponse(_PREVIEW_CSS + body, headers=_PREVIEW_HEADERS)


@app.get("/documents/{doc_id}/preview", tags=["documents"])
def document_preview(
    doc_id: uuid.UUID, sheet: int = 0, db: Session = Depends(get_session)
):
    """HTML preview of the document — PDF served directly; xlsx/xls converted to an HTML table."""
    doc = db.get(PriceDocument, doc_id)
    if doc is None:
        raise HTTPException(404, "document not found")
    path = _safe_original_path(doc)

    fmt = doc.file_format.lower()

    if fmt == "pdf":
        return FileResponse(
            str(path),
            filename=doc.file_name,
            media_type="application/pdf",
            content_disposition_type="inline",
            headers={"X-Content-Type-Options": "nosniff"},
        )

    if fmt in ("xlsx", "xls"):
        import html as html_mod

        import pandas as pd

        try:
            xls = pd.ExcelFile(str(path))
            sheet_name = xls.sheet_names[min(sheet, len(xls.sheet_names) - 1)]
            df = pd.read_excel(xls, sheet_name=sheet_name, header=0, dtype=str).fillna(
                ""
            )
            # pandas to_html escapes cell contents by default (escape=True)
            table_html = df.to_html(index=False, border=0, classes="data")
            sheet_tabs = "".join(
                f'<a href="?sheet={i}" style="margin-right:8px;font-weight:{"bold" if i == sheet else "normal"}">{html_mod.escape(str(s))}</a>'
                for i, s in enumerate(xls.sheet_names)
            )
            body = f"<div style='padding-bottom:8px;border-bottom:1px solid #ddd;margin-bottom:8px'>{sheet_tabs}</div>{table_html}"
        except Exception as exc:
            body = f"<p style='color:red'>Не удалось прочитать файл: {html_mod.escape(str(exc))}</p>"
        return _preview_html(body)

    if fmt in ("docx", "doc"):
        import html as html_mod

        try:
            from docx import Document as DocxDocument

            document = DocxDocument(str(path))
            parts: list[str] = []
            for para in document.paragraphs:
                if para.text.strip():
                    tag = "h3" if para.style.name.startswith("Heading") else "p"
                    parts.append(f"<{tag}>{html_mod.escape(para.text)}</{tag}>")
            for table in document.tables:
                rows_html = ""
                for i, row in enumerate(table.rows):
                    tag = "th" if i == 0 else "td"
                    cells = "".join(
                        f"<{tag}>{html_mod.escape(c.text)}</{tag}>" for c in row.cells
                    )
                    rows_html += f"<tr>{cells}</tr>"
                parts.append(f"<table>{rows_html}</table>")
            body = "\n".join(parts) if parts else "<p>Документ пуст</p>"
        except Exception as exc:
            body = f"<p style='color:red'>Не удалось прочитать документ: {html_mod.escape(str(exc))}</p>"
        return _preview_html(body)

    raise HTTPException(415, f"preview not supported for format: {fmt}")


# ----------------------------------------------------------- async ingestion jobs
@dataclass
class _Job:
    """In-memory state for one background ingestion (a zip bundle or a single file)."""

    id: str
    name: str
    status: str = "queued"  # queued | running | done | error
    total: int = 0
    done: int = 0
    current: str | None = None
    reports: list[dict] = field(default_factory=list)
    error: str | None = None
    started: float = 0.0
    finished: float | None = None

    def view(self) -> dict:
        reports = list(self.reports)  # snapshot — the worker thread may be appending
        return {
            "job_id": self.id,
            "archive": self.name,
            "status": self.status,
            "total": self.total,
            "done": self.done,
            "current": self.current,
            "files": len(reports),
            "errors": sum(r.get("status") == "error" for r in reports),
            "rows_written": sum(r.get("rows_written", 0) for r in reports),
            "rows_needs_review": sum(r.get("rows_needs_review", 0) for r in reports),
            "elapsed_sec": round((self.finished or time.time()) - self.started, 1),
            "error": self.error,
            "reports": reports,
        }


# Jobs live in-process (lost on restart) — fine for an interactive ingest console.
_JOBS: dict[str, _Job] = {}


def _run_job(
    job: _Job, tmp_path: Path, suffix: str, name: str, force: bool, ocr: bool
) -> None:
    """Background ingestion (Starlette runs this in a threadpool). Owns its own DB session
    so it's independent of any request lifecycle, and updates `job` for the poller."""
    from .db import SessionLocal

    job.status = "running"
    try:
        with SessionLocal() as session:
            if suffix == ".zip":

                def _progress(rep, done, total):
                    job.total = total
                    if rep is not None:
                        job.done = done
                        job.current = rep.file_name
                        job.reports.append(rep.as_dict())

                ingest_zip(
                    session, tmp_path, force=force, force_ocr=ocr, on_file=_progress
                )
            else:
                job.total = 1
                parseable = ensure_parseable(tmp_path)
                try:
                    rep = ingest_file(
                        session,
                        parseable,
                        original_name=name,
                        force=force,
                        force_ocr=ocr,
                    )
                finally:
                    if parseable != tmp_path:
                        Path(parseable).unlink(missing_ok=True)
                job.done = 1
                job.current = rep.file_name
                job.reports.append(rep.as_dict())
        job.status = "done"
    except Exception as e:  # noqa: BLE001
        job.status = "error"
        job.error = f"{type(e).__name__}: {e}"
    finally:
        Path(tmp_path).unlink(missing_ok=True)
        job.finished = time.time()


@app.post("/upload", tags=["documents"])
async def upload(
    background: BackgroundTasks,
    file: UploadFile = File(...),
    force: bool = False,
    ocr: bool = False,
):
    """Accept a .zip bundle (or a single .pdf/.docx/.doc/.xlsx/.xls) and ingest it in the
    BACKGROUND, returning a job id immediately. Poll GET /jobs/{job_id} for live progress
    and the per-file reports — the request never blocks on parsing/OCR."""
    name = file.filename or "upload"
    suffix = Path(name).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES | {".zip"}:
        allowed = ", ".join(sorted(SUPPORTED_SUFFIXES | {".zip"}))
        raise HTTPException(
            415, f"unsupported file type: {suffix} (allowed: {allowed})"
        )
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = Path(tmp.name)

    job = _Job(id=uuid.uuid4().hex[:12], name=name, started=time.time())
    _JOBS[job.id] = job
    background.add_task(_run_job, job, tmp_path, suffix, name, force, ocr)
    return {"job_id": job.id, "status_url": f"/jobs/{job.id}", "status": job.status}


@app.get("/jobs/{job_id}", tags=["documents"])
def job_status(job_id: str):
    """Live status + per-file reports for a background ingestion started via POST /upload."""
    job = _JOBS.get(job_id)
    if job is None:
        raise HTTPException(404, "unknown job id")
    return job.view()


@app.get("/stats", tags=["dashboard"])
def stats(db: Session = Depends(get_session)):
    items = db.scalar(select(func.count()).select_from(PriceItem)) or 0
    matched = (
        db.scalar(
            select(func.count())
            .select_from(PriceItem)
            .where(PriceItem.service_id.isnot(None))
        )
        or 0
    )
    review = (
        db.scalar(
            select(func.count())
            .select_from(PriceItem)
            .where(PriceItem.needs_review.is_(True))
        )
        or 0
    )
    by_status = dict(
        db.execute(
            select(PriceDocument.parse_status, func.count()).group_by(
                PriceDocument.parse_status
            )
        ).all()
    )
    return {
        "documents": db.scalar(select(func.count()).select_from(PriceDocument)) or 0,
        "documents_by_status": by_status,
        "partners": db.scalar(select(func.count()).select_from(Partner)) or 0,
        "services": db.scalar(select(func.count()).select_from(Service)) or 0,
        "price_items": items,
        "auto_matched": matched,
        "match_rate_pct": round(matched / items * 100, 1) if items else 0.0,
        "needs_review": review,
    }


def _partner_price(db: Session, it: PriceItem) -> PartnerPriceOut:
    return PartnerPriceOut(
        partner=PartnerOut.model_validate(db.get(Partner, it.partner_id)),
        price_resident_kzt=float(it.price_resident_kzt)
        if it.price_resident_kzt is not None
        else None,
        price_nonresident_kzt=float(it.price_nonresident_kzt)
        if it.price_nonresident_kzt is not None
        else None,
        currency_original=it.currency_original,
        effective_date=it.effective_date,
        service_name_raw=it.service_name_raw,
        item_id=it.item_id,
    )


# Static upload UI at "/".
if STATIC_DIR.exists():
    app.mount("/ui", StaticFiles(directory=str(STATIC_DIR), html=True), name="ui")


@app.get("/", include_in_schema=False)
def index():
    idx = STATIC_DIR / "index.html"
    if idx.exists():
        return FileResponse(idx)
    return {"service": "MedArchive", "docs": "/docs"}
